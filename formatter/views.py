from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from django.views.decorators.http import require_http_methods
from django.core.paginator import Paginator
from django.db.models import Q
from .utils import markdown_to_html
from .models import TextHistory
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
from xhtml2pdf import pisa
from io import BytesIO
import re
import logging

logger = logging.getLogger(__name__)

def instructions(request):
    """Instructions page view"""
    return render(request, "formatter/instructions.html")

def filter_history(request):
    """Filter history items by search query"""
    query = request.GET.get('q', '').strip()[:200]  # Limit query length
    
    if query:
        history_items = TextHistory.objects.filter(
            Q(raw_text__icontains=query) | Q(formatted_html__icontains=query)
        ).order_by('-created_at')[:50]
    else:
        history_items = TextHistory.objects.order_by('-created_at')[:10]

    return render(request, 'formatter/dashboard.html', {
        'history_items': history_items,
        'raw_text': '',
        'html_output': '',
        'search_query': query
    })

@require_http_methods(["POST"])
def delete_history(request, pk):
    """Delete a single history item"""
    try:
        TextHistory.objects.filter(pk=pk).delete()
        return JsonResponse({"status": "ok"})
    except Exception as e:
        logger.error(f"Error deleting history {pk}: {e}")
        return JsonResponse({"status": "error", "message": str(e)}, status=500)

@require_http_methods(["POST"])
def delete_all_history(request):
    """Delete all history items"""
    try:
        count = TextHistory.objects.all().delete()[0]
        return JsonResponse({"status": "ok", "deleted": count})
    except Exception as e:
        logger.error(f"Error deleting all history: {e}")
        return JsonResponse({"status": "error", "message": str(e)}, status=500)

def history_page(request):
    """Paginated history API endpoint"""
    page = int(request.GET.get('page', 1))
    items = TextHistory.objects.order_by('-created_at')
    paginator = Paginator(items, 10)
    page_obj = paginator.get_page(page)
    
    data = [{
        "id": i.id,
        "raw_text": i.raw_text[:100],
        "created_at": i.created_at.strftime("%Y-%m-%d %H:%M")
    } for i in page_obj]
    
    return JsonResponse({
        "items": data,
        "has_next": page_obj.has_next(),
        "has_previous": page_obj.has_previous()
    })

def dashboard(request):
    """Main dashboard view"""
    raw_text = ""
    html_output = ""
    error_message = None

    if request.method == "POST":
        raw_text = request.POST.get("raw_text", "").strip()
        
        # Check if this is a format request
        if "format" in request.POST:
            if not raw_text:
                error_message = "Please enter some text to format."
            else:
                try:
                    html_output = markdown_to_html(raw_text)
                    
                    # Save to history
                    TextHistory.objects.create(
                        raw_text=raw_text,
                        formatted_html=html_output
                    )
                except Exception as e:
                    logger.error(f"Error formatting text: {e}")
                    error_message = "An error occurred while processing your text."
        
        # Check if this is an export request
        elif "export_word" in request.POST or "export_pdf" in request.POST:
            html_output = request.POST.get("formatted_html", "")
            
            if not html_output:
                error_message = "No content to export. Please format text first."
            else:
                try:
                    filename = request.POST.get("file_name", "").strip()
                    
                    if "export_word" in request.POST:
                        return export_word(html_output, filename or "document.docx")
                    
                    if "export_pdf" in request.POST:
                        return export_pdf(html_output, filename or "document.pdf")
                        
                except Exception as e:
                    logger.error(f"Error exporting: {e}")
                    error_message = f"An error occurred while exporting: {str(e)}"

    # Load history
    history_items = TextHistory.objects.order_by('-created_at')[:10]

    return render(request, "formatter/dashboard.html", {
        "raw_text": raw_text,
        "html_output": html_output,
        "history_items": history_items,
        "error_message": error_message
    })

def load_history(request, pk):
    """Load a specific history item"""
    try:
        item = TextHistory.objects.get(pk=pk)
        return JsonResponse({
            "raw_text": item.raw_text,
            "formatted_html": item.formatted_html
        })
    except TextHistory.DoesNotExist:
        return JsonResponse({"error": "Item not found"}, status=404)
    except Exception as e:
        logger.error(f"Error loading history {pk}: {e}")
        return JsonResponse({"error": "Server error"}, status=500)

def sanitize_filename(name):
    """Remove unsafe characters from filename"""
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    name = name[:200]  # Limit length
    return name or "document"

def export_word(html, filename="document.docx"):
    """Export formatted HTML to Word document"""
    try:
        logger.info(f"Exporting Word document: {filename}")
        logger.info(f"HTML content length: {len(html)}")
        
        filename = sanitize_filename(filename)
        if not filename.lower().endswith(".docx"):
            filename += ".docx"
            
        soup = BeautifulSoup(html, "html.parser")
        doc = Document()

        # Parse all elements
        for element in soup.children:
            parse_html_element(doc, element)

        # If document is empty, add a message
        if len(doc.paragraphs) == 0:
            doc.add_paragraph("No content available")
            logger.warning("Generated empty Word document")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        logger.info(f"Word document exported successfully: {filename}")
        return response
        
    except Exception as e:
        logger.error(f"Error exporting Word document: {e}", exc_info=True)
        raise

def parse_html_element(doc, element):
    """Parse HTML element and add to Word document"""
    if not hasattr(element, "name"):
        return

    if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
        level = min(int(element.name[1]), 9)
        doc.add_heading(element.get_text(), level=level)

    elif element.name == "p":
        p = doc.add_paragraph()
        add_inline_formatting(p, element)

    elif element.name == "ul":
        for li in element.find_all("li", recursive=False):
            doc.add_paragraph(li.get_text(), style="List Bullet")

    elif element.name == "ol":
        for li in element.find_all("li", recursive=False):
            doc.add_paragraph(li.get_text(), style="List Number")

    elif element.name in ["pre", "code"]:
        run = doc.add_paragraph().add_run(element.get_text())
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    
    elif element.name == "blockquote":
        p = doc.add_paragraph(element.get_text())
        p.style = "Quote"

def add_inline_formatting(paragraph, element):
    """Add inline formatting (bold, italic, code) to paragraph"""
    for node in element.contents:
        if isinstance(node, str):
            paragraph.add_run(node)
        elif hasattr(node, 'name'):
            if node.name == "strong" or node.name == "b":
                paragraph.add_run(node.get_text()).bold = True
            elif node.name == "em" or node.name == "i":
                paragraph.add_run(node.get_text()).italic = True
            elif node.name == "code":
                run = paragraph.add_run(node.get_text())
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            elif node.name == "a":
                run = paragraph.add_run(node.get_text())
                run.font.color.rgb = None  # Blue hyperlink color
            else:
                paragraph.add_run(node.get_text())

def export_pdf(html, filename="document.pdf"):
    """Export formatted HTML to PDF"""
    try:
        logger.info(f"Exporting PDF document: {filename}")
        logger.info(f"HTML content length: {len(html)}")
        
        filename = sanitize_filename(filename)
        if not filename.lower().endswith(".pdf"):
            filename += ".pdf"
        
        # If HTML is empty, add message
        if not html or len(html.strip()) == 0:
            html = "<p>No content available</p>"
            logger.warning("Generating PDF with empty content")
            
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: Arial, sans-serif;
                    font-size: 12px;
                    line-height: 1.6;
                    color: #333;
                }}
                h1, h2, h3, h4, h5, h6 {{
                    color: #2c3e50;
                    margin-top: 1em;
                    margin-bottom: 0.5em;
                }}
                h1 {{ font-size: 24px; }}
                h2 {{ font-size: 20px; }}
                h3 {{ font-size: 16px; }}
                p {{ margin-bottom: 1em; }}
                code {{
                    background-color: #f4f4f4;
                    padding: 2px 5px;
                    font-family: 'Courier New', monospace;
                    font-size: 11px;
                }}
                pre {{
                    background-color: #f4f4f4;
                    padding: 10px;
                    border-left: 3px solid #ccc;
                    overflow-x: auto;
                }}
                ul, ol {{ margin-left: 20px; }}
                blockquote {{
                    border-left: 4px solid #ddd;
                    padding-left: 15px;
                    color: #666;
                    font-style: italic;
                }}
            </style>
        </head>
        <body>
            {html}
        </body>
        </html>
        """

        buffer = BytesIO()
        pisa_status = pisa.CreatePDF(full_html, dest=buffer)
        
        if pisa_status.err:
            logger.error("PDF generation failed with errors")
            raise Exception("PDF generation failed")

        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        logger.info(f"PDF document exported successfully: {filename}")
        return response
        
    except Exception as e:
        logger.error(f"Error exporting PDF: {e}", exc_info=True)
        raise